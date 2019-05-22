from docx import Document

# Creating a word file object
doc = open('data/file.docx', 'rb')

# creating word reader object
document = Document(doc)

# create an empty string and call this document. This document
# variable store each paragraph in the Word document.We then
# create a for loop that goes through each paragraph in the Word
# document and appends the paragraph.

docu = ""

for paragraph in document.paragraphs:
    docu += paragraph.text

# To see the output call docu
print(docu)

# doc.close()

# or 

with open('data/file.docx', 'rb') as word:
    document = Document(doc)
    docu = ""

    for paragraph in document.paragraphs:
        docu += paragraph.text

    # To see the output call docu
    print(docu)
